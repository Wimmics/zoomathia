import unicodedata
import re
from glob import glob
from os import getcwd
from time import time
import difflib

# Corese execute
import atexit
import subprocess
from time import sleep
from py4j.java_gateway import JavaGateway
from py4j.java_collections import JavaMap

# Word manipulation
import win32com.client as win32
from os import path

import pandas as pd
from docx import Document
import zipfile
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


# Start java gateway
java_process = subprocess.Popen(
    ['java', '-jar', '-Dfile.encoding=UTF-8', 'corese-library-python-4.4.1.jar'])
sleep(1)
gateway = JavaGateway()

# Stop java gateway at the enf od script
def exit_handler():
    gateway.shutdown()
    print('\n' * 2)
    print('Gateway Server Stop!')

atexit.register(exit_handler)

# Import of class
Graph = gateway.jvm.fr.inria.corese.core.Graph
Load = gateway.jvm.fr.inria.corese.core.load.Load
Transformer = gateway.jvm.fr.inria.corese.core.transform.Transformer
QueryProcess = gateway.jvm.fr.inria.corese.core.query.QueryProcess
CoreseModel = gateway.jvm.fr.inria.corese.rdf4j.CoreseModel
SimpleValueFactory = gateway.jvm.org.eclipse.rdf4j.model.impl.SimpleValueFactory
RDF = gateway.jvm.org.eclipse.rdf4j.model.vocabulary.RDF

name_file = ""


nb_commentaire, nb_concepts, nb_paragraph = 0, 0, 0

def get_full_extraction(filepath, chapter):
    doc = word.Documents.Open(filepath)

    comments = doc.Comments
    # Parcourir les commentaires
    extractions_full = []
    id_annotation = 0
    for comment in comments:
        # Récupérer l'objet Range représentant le commentaire
        scope_comment = comment.Scope
        # Récupérer l'objet Range correspondant au paragraphe contenant le commentaire
        mention = re.sub(r"\[\[([0-9]*)\]\]", "", scope_comment.Text.encode("utf-8").decode()).strip()\
            .replace("\n", "").replace("\r", "").replace("<", "").replace(">", "").replace(u'\u8232', "").replace(u"\xa0", "").replace(u"\x0b", "")
        # Récupère le text complet du paragraph
        # paragraph_complet = scope_comment.Paragraphs(1)
        # Récupère la portion de texte du paragraphe commentée
        comment_text = unicodedata.normalize("NFKD", comment.Range.Text.encode("utf-8").decode().strip().replace("â€™", "'"))

        for p in scope_comment.Paragraphs:
            id_paragraph = re.search(r"\[\[([0-9]*)\]\]", p.Range.Text).group(1)
            paragraph_clean = re.sub(r"\[\[([0-9]*)\]\]", "", p.Range.Text).strip()\
                .replace("\n", "").replace("\r", "").replace("<", "").replace(">", "").replace(u'\u8232', "").replace(u"\xa0", "").replace(u"\x0b", "")

            matcher = difflib.SequenceMatcher(None, paragraph_clean, mention, autojunk=False)
            match = matcher.find_longest_match(0, len(paragraph_clean), 0, len(mention))
            if match.size > 0:
                start = match.a
                end = match.a + match.size
                size = match.size
                common_match = paragraph_clean[start:end]
            else:
                # Error : la mention ne correspond pas aux paragraphes
                print(f"Aucun matching pour {mention}")

            """["annotation_id", "paragraphe_id", "paragraph_text",
             "chapter",
             "concepts", "mention", "common_text", "start", "end", "size"]"""
            extractions_full.append(
                [id_annotation, id_paragraph, paragraph_clean,
                 chapter,
                 comment_text.strip(), mention.strip(), common_match, start, end, size]
            )
        id_annotation += 1
    return extractions_full


def sparqlQuery(graph, query):
    """Run a query on a graph

    :param graph: the graph on which the query is executed
    :param query: query to run
    :returns: query result
    """
    exec = QueryProcess.create(graph)
    return exec.query(query)


def load(path):
    """Load a graph from a local file or a URL

    :param path: local path or a URL
    :returns: the graph load
    """
    graph = Graph()

    ld = Load.create(graph)
    ld.parse(path)

    return graph


def error_detect_pattern(label_concept, row_obj):
    id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention, common_text, start, end, size = row_obj

    if label_concept.isdigit() or "year" in label_concept:
        error_found.append(["Numeric_Value", label_concept,
                            id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                            common_text, start, end, size
                            ])
    elif isIncompleteLabel(label_concept):
        error_found.append(["IncompleteLabel", label_concept,
                            id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                            common_text, start, end, size
                            ])
    else:
        error_found.append(["No_Concept", label_concept,
                            id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                            common_text, start, end, size
                            ])


def concept_search(annoted_list):
    data_list = []
    for row in annoted_list:
        id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention, common_text, start, end, size = row
        distinct_concept = concepts.split(";")
        for distinct in distinct_concept:
            spread_concept = list(map(lambda x: x.replace("\r", "").replace("\n", "").replace(u'\u8232', "").replace(u"\xa0", "").strip(), distinct.split(":")))
            parent = ""
            for label_concept in spread_concept:
                if label_concept == "":
                    continue
                for individual_label in label_concept.split(","):
                    if "?" in individual_label:
                        toCheck.append(["Need_validation", individual_label,
                                            id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                                            common_text, start, end, size
                                            ])
                    label_clean = individual_label.encode("utf-8").decode().strip().replace("?", "")
                    result = skos_search(label_clean, parent, row)
                    if len(result) > 1:
                        """["error_type", "error_label", "annotation_id", "paragraph_id", "paragraph_text",
                            "chapter",
                            "concepts", "mention", "common_text", "start", "end", "size"]"""
                        error_found.append([ "Ambigue", label_clean,
                             id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                             common_text, start, end, size
                        ])
                    elif result == []:
                        error_detect_pattern(label_clean, [id_annotation, paragraph_id, paragraph_text, chapter, concepts, mention,
                             common_text, start, end, size
                        ])
                    else:
                        concept_id = result[0].split("?idg=")[-1].split("?idc=")[-1].split("&idt=")[0]
                        if "?idg=" in result[0]:
                            data_list.append([id_annotation, "Collection", concept_id, paragraph_id, paragraph_text, chapter, concepts, mention, common_text, start, end, size])
                        else:
                            data_list.append([id_annotation, "Concept", concept_id, paragraph_id, paragraph_text, chapter, concepts, mention, common_text, start, end, size])
                parent = label_concept
    return data_list


def isIncompleteLabel(label):
    # check if label is contain in prefLabel
    # Set incompleteLabel then
    q = f"""prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

    SELECT DISTINCT * WHERE {{
        ?x a ?type;
            skos:prefLabel ?label.
        FILTER(lang(?label) = "en" || lang(?label) = "fr").
        FILTER (contains(str(?label), "{label}")).
    }}
    """

    query_result = sparqlQuery(graph, q)
    candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
    return candidate


def skos_search(label, parent, row_obj):
    if parent != "":
        q = f"""
    prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

select * where {{
    {{
        ?x skos:prefLabel ?label;
            a ?type.
                  ?y skos:prefLabel ?collection;
        skos:member ?x.
        filter(lang(?collection) = "en" || lang(?collection) = "fr").
        FILTER(lang(?label) = "en" || lang(?label) = "fr").
        filter("{label}" in (ucase(str(?label)), lcase(str(?label)), str(?label))).
        filter( "{parent}" in (ucase(str(?collection)), lcase(str(?collection)), str(?collection))).
  }} UNION {{
        ?x skos:prefLabel ?label;
                a ?type;
                skos:broader+ ?y.
                ?y skos:prefLabel ?concept;
        filter(lang(?concept) = "en" || lang(?concept) = "fr").
        FILTER(lang(?label) = "en" || lang(?label) = "fr").
        filter("{label}" in (ucase(str(?label)), lcase(str(?label)), str(?label))).
        filter( "{parent}" in (ucase(str(?concept)), lcase(str(?concept)), str(?concept))).
  }}
}}
"""
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))

        if candidate == []:
            # Empty = Hierachy broken, no result for this descendant
            # len > 1 = Ambigus respons for this concept
            return skos_search(label, "", row_obj)
        if len(candidate) > 1 :
            # Ambigus respons for label with this hierarchy
            return skos_search(label, "", row_obj)
        return candidate
    else:
        q = f"""
    prefix skos: <http://www.w3.org/2004/02/skos/core#>  .
    
    SELECT DISTINCT * WHERE {{
        ?x a ?type;
            skos:prefLabel ?label.
        FILTER(lang(?label) = "en" || lang(?label) = "fr").
        FILTER ("{label}" in (ucase(str(?label)), lcase(str(?label)), str(?label))).
    }}
    """
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
        if candidate == []:
            # No concept found for isolated label
            return candidate
        if len(candidate) > 1:
            # Ambigus concept for isolated label
            return candidate
        return candidate


if __name__ == '__main__':

    alpha = time()
    no_concept_found = list()
    many_candidates = list()

    graph = load("th310.ttl")

    files = [x for x in glob("PLINE-*.docx")]

    # Lecture du fichier word: lance un word en arrière plan.

    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False

    annotation_list = []
    annotation_result = []
    error_found = []
    paragraph_list = []
    toCheck = []

    for file in files:
        print(file)
        name_file = file
        filepath = path.normpath(f"{getcwd()}\\{file}")

        doc_obj = Document(filepath)
        docxZip = zipfile.ZipFile(filepath)

        para_content = [p.text for p in doc_obj.paragraphs]
        nb_paragraph += len(doc_obj.paragraphs)
        paragraph_list.extend([[
            f"{file.split('.')[0]}",
            para_content[0].split("]]")[-1].strip(),
            para_content[1].split("]]")[-1].strip().replace("Historia naturalis, ", ""),
            para_content[2].split("]]")[-1].strip(),
            file.split(".")[0].split("-")[1],
            "".join(p.split("]]")[0].split("[[")[-1]).strip(),
            "".join(p.encode("utf-8").decode().split("]]")[-1]).strip().replace("<", "").replace(">", "").replace(u'\u8232', "").replace(u"\xa0", "").replace(u"\x0b", ""),
            "Historia naturalis",
        "historia_naturalis"]
            for p in para_content[3:] if p != ""])

        annotation_list.extend(get_full_extraction(filepath, file.split(".")[0].split("-")[1]))

    """treatment = pd.DataFrame(annotation_list, columns=["annotation_id", "paragraphe_id", "paragraph_text",
             "chapter",
             "concepts", "mention", "common_text", "start", "end", "size"])
    treatment.to_csv("treatment.csv", index=False, encoding="utf-8", sep=";", lineterminator="\n")"""
    annotation_result.extend(concept_search(annotation_list))

    no_concept = pd.DataFrame(error_found, columns=["error_type", "error_label", "annotation_id", "paragraph_id", "paragraph_text",
                        "chapter",
                        "concepts", "mention", "common_text", "start", "end", "size"])
    to_check = pd.DataFrame(toCheck, columns=["error_type", "error_label", "annotation_id", "paragraph_id", "paragraph_text",
                        "chapter",
                        "concepts", "mention", "common_text", "start", "end", "size"])
    to_check.to_csv(f"to_check.csv", index=False, encoding="utf-8", sep=";", lineterminator="\n")
    no_concept.to_csv(f"error_found_new.csv", index=False, encoding="utf-8", sep=";", lineterminator="\n")

    extract_para = pd.DataFrame(paragraph_list,
                                columns=["name", "title", "author", "edition", "chapter", "paragraph_number",
                                         "paragraphe_text", "oeuvre", "uri"])
    extract_para.to_csv(f"paragraphs.csv", index=False, encoding="utf-8", sep=";", lineterminator="\n")
    extraction_annotation = pd.DataFrame(annotation_result, columns=["annotation_id", "concept_type", "concept_id", "paragraph_id", "paragraph_text",
                                                                   "chapter",
                                                                   "concepts", "mention", "common_text", "start", "end", "size"])

    extraction_annotation.to_csv(f"comment_extraction.csv", index=False, encoding="utf-8", sep=";", lineterminator="\n")

    print(f"Traitement fini en '{time() - alpha}' secondes")
    print(f"nombre de paragraphes {nb_paragraph}, nombre de commentaires {nb_commentaire}, nombre de concepts: {nb_concepts}")
